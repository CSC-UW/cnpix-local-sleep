"""Generate SLEEP-journal-ready statistical tables for the manuscript.

This builds on :mod:`summarize_results` (whose ``results.json`` readers and
contrast-matching are reused) but targets *publication* output rather than the
quick markdown/CSV summaries used for internal review:

* SLEEP formatting: <= 2 significant digits; adjusted p with significance stars;
  Cohen's-d analogue and its 95% CI in every post-hoc cell; Cohen's f2 in the
  main-effect cell.
* Human property labels + units pulled from the ``title`` field of the analysis
  config (``config/cx_homeostasis.yaml``), and human partition labels
  (All OFFs / Medium + Large / Small OFFs) from the summary config.
* Manuscript scope only: the summary config lists the exact entries to report, so
  BLAS (which never appears in the manuscript) is simply never listed.
* Editable .docx (Word) + .xlsx (Excel) emitters, plus the CSV that the
  older ``summarize_results.py`` already produced, for quick diffs.

Each "S-table" is described by a YAML config under ``config/summary_tables/``.
The homeostasis table (S1a) reuses the same ``entries`` list schema as
``summary.yaml`` (``[response_variable, dataset, condition_set, model]``) plus a
``partition_labels`` map. See ``config/summary_tables/manuscript_s1a_homeostasis.yaml``.

Usage:
    python manuscript_tables.py config/summary_tables/manuscript_s1a_homeostasis.yaml
    python manuscript_tables.py config/summary_tables/manuscript_s1a_homeostasis.yaml \
        --docx out/S1a.docx --xlsx out/S1a.xlsx
"""

from __future__ import annotations

import argparse
import json
import math
import sys
from collections import OrderedDict
from pathlib import Path

import yaml

# Reuse the proven readers/contrast-matching from summarize_results.py (SPOT).
sys.path.insert(0, str(Path(__file__).resolve().parent))
from summarize_results import (  # noqa: E402
    get_condition_set_posthocs,
    load_results,
    normalize_contrast,
)

# Partition column order for the homeostasis table (All = union; Medium+Large =
# stringent subset = CLAS; Small = relaxed-not-stringent = llas & ~clas).
PARTITION_ORDER = ["llas", "clas", "llas_exclusive"]
# Condition-set order within a property block (aggregate measures are fit as
# independent NREM and Wake models).
SET_ORDER = {"six": 0, "nrem": 1, "wake": 2}


# SLEEP-compliant cell formatting (<= 2 significant digits)
def _stars(p: float) -> str:
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return ""


def _sig2(x: float | None) -> str:
    """Format a number to 2 significant digits (SLEEP: '<= 2 relevant digits'),
    avoiding scientific notation for values in a normal reading range."""
    if x is None:
        return ""
    if x == 0 or not math.isfinite(x):
        return "0" if x == 0 else "n/a"
    if abs(x) < 1e-3 or abs(x) >= 1e5:
        return f"{x:.1e}"
    ndig = 1 - int(math.floor(math.log10(abs(x))))  # decimals for 2 sig figs
    if ndig <= 0:
        return f"{round(x, ndig):.0f}"
    return f"{x:.{ndig}f}"


def _at(lst, idx):
    """Safe list index: None if missing or out of range."""
    if not lst or idx >= len(lst):
        return None
    return lst[idx]


def fmt_p(p: float) -> str:
    """p-value string, thresholded at 0.001, else 2 significant digits."""
    if p < 0.001:
        return "p < 0.001"
    return f"p = {p:.2g}"


def fmt_p_starred(p: float) -> str:
    s = _stars(p)
    return f"{fmt_p(p)} {s}".rstrip()


def fmt_main_effect(main_effect: dict) -> str:
    """Main-effect cell: adjusted omnibus p (+ stars) and Cohen's f2 if significant."""
    p = main_effect["pval"]
    cell = fmt_p_starred(p)
    if main_effect.get("significant") and main_effect.get("cohens_f2") is not None:
        cell += f" (f² = {_sig2(main_effect['cohens_f2'])})"
    return cell


def fmt_est_ci_d(posthoc: dict, idx: int) -> str:
    """'estimate [lo, hi], d = X': the raw contrast estimate with its 95% CI (on
    the response scale, units given in the table's Units column) and the Cohen's-d
    analogue as a unitless standardized effect size."""
    parts = []
    est = _at(posthoc.get("estimates"), idx)
    if est is not None:
        s = _sig2(est)
        lo, hi = _at(posthoc.get("ci_lower"), idx), _at(posthoc.get("ci_upper"), idx)
        if lo is not None and hi is not None:
            s += f" [{_sig2(lo)}, {_sig2(hi)}]"
        parts.append(s)
    d = _at(posthoc.get("cohens_d"), idx)
    if d is not None:
        parts.append(f"d = {_sig2(d)}")
    return ", ".join(parts)


def fmt_contrast(posthoc: dict, idx: int) -> str:
    """Post-hoc cell: adjusted p (+ stars), then the contrast estimate + 95% CI + d."""
    body = fmt_est_ci_d(posthoc, idx)
    p = fmt_p_starred(_at(posthoc.get("pvalues"), idx))
    return f"{p}, {body}" if body else p


# -------------------- Config plumbing --------------------
def load_rv_titles(main_config: dict) -> dict:
    """Map (dataset, response_variable) -> human title (with units) from the
    analysis config's ``datasets.<ds>.response_variables[].title`` fields.
    Falls back across datasets so a partition can borrow another's title.
    """
    titles: dict = {}
    generic: dict = {}
    for ds, ds_cfg in (main_config.get("datasets") or {}).items():
        for rv in ds_cfg.get("response_variables") or []:
            name = rv.get("response_variable")
            title = rv.get("title")
            if name and title:
                titles[(ds, name)] = title
                generic.setdefault(name, title)
    titles["_generic"] = generic
    return titles


def resolve_base_dir(summary_yaml: Path, summary_config: dict) -> Path:
    base = summary_yaml.resolve().parent
    results_dir_name = summary_config.get("results_dir", "_output")
    while base != base.parent:
        if (base / results_dir_name).is_dir():
            return base
        base = base.parent
    raise SystemExit(f"Could not find '{results_dir_name}' above {summary_yaml}")


# Table construction (homeostasis-style: property blocks x partition columns)
class Cell:
    """A rendered cell plus whether it is (statistically) significant, so the
    emitters can optionally bold/mark it."""

    __slots__ = ("text", "significant")

    def __init__(self, text: str = "", significant: bool = False):
        self.text = text
        self.significant = significant

    def __repr__(self):  # pragma: no cover - debugging aid
        return f"Cell({self.text!r})"


def _section_configs(summary_config: dict) -> list:
    """Normalize the config into a list of section dicts. A config may either
    carry a top-level ``entries`` (one implicit, untitled section) or a
    ``sections:`` list (each with its own entries and optional results_dir/config/
    labels overriding the top-level defaults)."""
    if summary_config.get("sections"):
        return summary_config["sections"]
    return [{"title": "", "entries": summary_config["entries"]}]


def _resolve(section: dict, summary_config: dict, key: str, default=None):
    return section.get(key, summary_config.get(key, default))


def build_section(section, summary_config, base_dir, all_contrasts):
    """Build the property blocks for one section against the *unified* contrast
    column order ``all_contrasts`` (norm -> label)."""
    results_dir = base_dir / _resolve(section, summary_config, "results_dir", "_output")
    main_config = yaml.safe_load(
        open(base_dir / _resolve(section, summary_config, "config"))
    )
    partition_labels = {
        "llas": "All OFFs", "clas": "Medium + Large", "llas_exclusive": "Small OFFs",
        **(summary_config.get("partition_labels") or {}),
        **(section.get("partition_labels") or {}),
    }
    set_labels = {
        "six": "—", "nrem": "NREM", "wake": "Wake",
        **(summary_config.get("set_labels") or {}),
        **(section.get("set_labels") or {}),
    }
    property_overrides = {
        **(summary_config.get("property_labels") or {}),
        **(section.get("property_labels") or {}),
    }
    property_units = {
        **(summary_config.get("property_units") or {}),
        **(section.get("property_units") or {}),
    }
    titles = load_rv_titles(main_config)

    def property_key(rv):
        return property_overrides.get(rv) or titles["_generic"].get(rv, rv)

    blocks: "OrderedDict[str, dict]" = OrderedDict()
    for rv, ds, cset, model in section["entries"]:
        block = blocks.setdefault(property_key(rv), {"label": property_key(rv), "rows": []})
        me = load_results(results_dir, ds, cset, rv, model)["main_effect"]
        cells = [
            Cell(property_units.get(rv, "")),
            Cell(partition_labels.get(ds, ds.upper())),
            Cell(set_labels.get(cset, cset)),
            Cell(fmt_main_effect(me), me.get("significant", False)),
        ]
        cs_norms = [normalize_contrast(s) for s in get_condition_set_posthocs(main_config, cset)]
        posthoc = me.get("posthoc") if me.get("significant") else None
        json_norms = (
            [normalize_contrast(c) for c in posthoc["contrasts"]]
            if posthoc and posthoc.get("contrasts") else None
        )
        for norm in all_contrasts:
            if norm not in cs_norms:
                cells.append(Cell("N/A"))
            elif posthoc is None:
                cells.append(Cell("—"))
            else:
                idx = (
                    json_norms.index(norm) if json_norms and norm in json_norms
                    else (cs_norms.index(norm) if norm in cs_norms else None)
                )
                if idx is None or idx >= len(posthoc.get("pvalues", [])):
                    cells.append(Cell("—"))
                else:
                    cells.append(Cell(fmt_contrast(posthoc, idx), posthoc["pvalues"][idx] < 0.05))
        sort = (PARTITION_ORDER.index(ds) if ds in PARTITION_ORDER else 99, SET_ORDER.get(cset, 9))
        block["rows"].append({"cells": cells, "sort": sort})

    for block in blocks.values():
        block["rows"].sort(key=lambda r: r["sort"])
        for r in block["rows"]:
            del r["sort"]
    return list(blocks.values())


def build_document(summary_config: dict, base_dir: Path):
    """Return (headers, sections) where sections = [{"title", "blocks"}]."""
    sections_cfg = _section_configs(summary_config)

    # Pass 1: unified contrast column order (union across all sections).
    all_contrasts: "OrderedDict[str, str]" = OrderedDict()
    contrast_labels = summary_config.get("contrast_labels", {})
    for section in sections_cfg:
        main_config = yaml.safe_load(
            open(base_dir / _resolve(section, summary_config, "config"))
        )
        section_labels = {**contrast_labels, **(section.get("contrast_labels") or {})}
        for _rv, _ds, cset, _model in section["entries"]:
            for ph in get_condition_set_posthocs(main_config, cset):
                norm = normalize_contrast(ph)
                if norm not in all_contrasts:
                    all_contrasts[norm] = section_labels.get(ph, ph)
    headers = ["Units", "Detection", "Model", "Main effect"] + list(all_contrasts.values())

    # Pass 2: build each section against the unified columns.
    out = []
    for section in sections_cfg:
        blocks = build_section(section, summary_config, base_dir, all_contrasts)
        out.append({"title": section.get("title", ""), "blocks": blocks})
    return ["Property"] + headers, out


def _ci_str(lo, hi):
    if lo is None or hi is None:
        return ""
    return f"[{_sig2(lo)}, {_sig2(hi)}]"


DEFAULT_PARTITIONS = {"llas": "All OFFs", "clas": "Medium + Large",
                      "llas_exclusive": "Small OFFs", "blas": "BLAS"}


def _correlation_entry(entry):
    """Normalize an S4 config entry to (dataset, metric, response_metric_or_None,
    predictor, response_metric_resolved)."""
    dataset, metric, response_metric, predictor = (list(entry) + [None] * 4)[:4]
    return dataset, metric, response_metric, predictor, (response_metric or metric)


def _read_csv_index(path: Path, key_cols: tuple) -> dict:
    """Read a tidy CSV into {tuple-of-key-columns: row-dict}."""
    import csv

    if not path.exists():
        raise FileNotFoundError(f"Required summary CSV not found: {path}")
    out = {}
    with open(path) as f:
        for row in csv.DictReader(f):
            out[tuple(row[c] for c in key_cols)] = row
    return out


XTRANSFORM_KEYS = ("predictor", "off_type", "metric", "response_metric", "model", "transform")


def _num(row: dict, key: str):
    """Float from a CSV cell, None for blanks / NA / non-numeric."""
    try:
        v = float(row.get(key, ""))
    except (TypeError, ValueError):
        return None
    return None if v != v else v  # drop NaN


def _assert_xtransform_current(raw_row: dict | None, base: dict, where: str):
    """Guard against pairing a stale predictor-transform sweep with freshly re-run
    base fits: the sweep's own ``raw`` refit must reproduce the base fit exactly."""
    if raw_row is None:
        raise KeyError(f"No 'raw' row in the transform sweep for {where}")
    for field in ("slope", "pval"):
        got, want = _num(raw_row, field), base[field]
        if got is None or not math.isclose(got, want, rel_tol=1e-6, abs_tol=1e-12):
            raise SystemExit(
                f"Stale xtransform_sensitivity_summary.csv: {where} raw {field} = {got}, "
                f"but the base fit has {want}. Re-run "
                "scripts/run_correlation_xtransform_sensitivity.R."
            )


def build_correlation_slope(summary_config: dict, base_dir: Path):
    """S4a: mixed-model slope that an SD-window OFF metric predicts its NREM
    rebound, with the model variants the manuscript promises (an added by-structure
    random intercept; rank- and log-transformed predictors).

    Raw-predictor rows come from the flat correlation ``results.json``
    (_output_correlations/<dataset>/<metric>[__vs__<resp>]/<pred>/<model>/).
    Transformed-predictor rows come from ``xtransform_sensitivity_summary.csv``,
    NOT the per-fit sensitivity ``results.json``, which the R writer rounds to 4 dp.

    Each config entry: [dataset, metric, response_metric|null, predictor]; the row
    set within an entry is the config's ``variants`` list ({model, transform, label}).
    """
    results_dir = base_dir / summary_config.get("results_dir", "_output_correlations")
    partition_labels = {**DEFAULT_PARTITIONS, **(summary_config.get("partition_labels") or {})}
    metric_labels = summary_config.get("metric_labels", {})
    predictor_labels = summary_config.get("predictor_labels", {})
    metric_units = summary_config.get("metric_units", {})
    variants = summary_config.get("variants") or [
        {"model": "subject", "transform": "raw", "label": "(1 | subject), raw predictor"}
    ]
    needs_xtab = any(v.get("transform", "raw") != "raw" for v in variants)
    xtab = (
        _read_csv_index(results_dir / "xtransform_sensitivity_summary.csv",
                        XTRANSFORM_KEYS)
        if needs_xtab else {}
    )
    headers = ["OFF detection / SD window", "Model variant", "n", "Slope [95% CI]",
               "Units", "p", "Cohen's f²"]
    sections: "OrderedDict[str, dict]" = OrderedDict()
    for entry in summary_config["entries"]:
        dataset, metric, response_metric, predictor, resp = _correlation_entry(entry)
        metric_dir = metric if not response_metric else f"{metric}__vs__{response_metric}"
        section = sections.setdefault(
            metric, {"title": metric_labels.get(metric, metric), "blocks": []})
        rows = []
        for v in variants:
            model, transform = v.get("model", "subject"), v.get("transform", "raw")
            path = (results_dir / dataset / metric_dir / predictor / model
                    / "results.json")
            if not path.exists():
                raise FileNotFoundError(f"Correlation results not found: {path}")
            with open(path) as f:
                base = json.load(f)
            where = f"{predictor}/{dataset}/{metric_dir}/{model}"
            if transform == "raw":
                slope, lo, hi = base["slope"], base.get("ci_lower"), base.get("ci_upper")
                pval, f2, n = base["pval"], base.get("cohens_f2"), base.get("n")
            else:
                key = (predictor, dataset, metric, resp, model, transform)
                row = xtab.get(key)
                if row is None:
                    raise KeyError(f"No {transform!r} row in the transform sweep for {where}")
                _assert_xtransform_current(
                    xtab.get((*key[:-1], "raw")), base, where)
                slope, lo, hi = _num(row, "slope"), _num(row, "ci_lower"), _num(row, "ci_upper")
                pval, f2, n = _num(row, "pval"), _num(row, "cohens_f2"), row.get("n", "")
            # f² is reported for every fit (the correlation pipeline always computes
            # it), not only significant ones; the Fig. 5 legend promises it for all.
            rows.append({"cells": [
                Cell(v["label"]),
                Cell(str(n or "")),
                Cell(f"{_sig2(slope)} {_ci_str(lo, hi)}".strip(), pval < 0.05),
                Cell((metric_units.get(metric) or {}).get(transform, "")),
                Cell(fmt_p_starred(pval), pval < 0.05),
                Cell(_sig2(f2) if f2 is not None else ""),
            ]})
        section["blocks"].append({
            "label": (f"{partition_labels.get(dataset, dataset.upper())}, "
                      f"{predictor_labels.get(predictor, predictor)}"),
            "rows": rows,
        })
    return headers, list(sections.values())


ROBUSTNESS_KEYS = ("predictor", "off_type", "metric", "response_metric", "model")


def _flip_cell(row: dict, flips_col: str, n_col: str) -> str:
    """'0 / 15': deletions that flip the p < 0.05 verdict, out of those performed.
    An em dash when the deletion does not apply (e.g. leaving out a structure in a
    model that has no structure grouping factor)."""
    n = _num(row, n_col)
    if not n:
        return "—"
    return f"{int(_num(row, flips_col) or 0)} / {int(n)}"


def build_correlation_robustness(summary_config: dict, base_dir: Path):
    """S4b: the leave-out / influence battery behind the S4a slopes, read from
    ``_output_correlations/influence_diagnostics/influence_summary.csv``
    (written by scripts/run_correlation_influence_diagnostics.R). Same entries and
    block structure as S4a, with one row per random-effects structure."""
    results_dir = base_dir / summary_config.get("results_dir", "_output_correlations")
    tab = _read_csv_index(
        results_dir / "influence_diagnostics" / "influence_summary.csv",
        ROBUSTNESS_KEYS)
    partition_labels = {**DEFAULT_PARTITIONS, **(summary_config.get("partition_labels") or {})}
    metric_labels = summary_config.get("metric_labels", {})
    predictor_labels = summary_config.get("predictor_labels", {})
    models = summary_config.get("models") or [
        {"model": "subject", "label": "(1 | subject)"},
        {"model": "subject_structure", "label": "+ (1 | structure)"},
    ]
    headers = ["OFF detection / SD window", "Model", "p", "Leave-one-subject-out",
               "Leave-one-structure-out", "Leave-one-pair-out", "Leave-two-subjects-out",
               "Leave-three-subjects-out", "Excluding two most influential subjects"]
    sections: "OrderedDict[str, dict]" = OrderedDict()
    for entry in summary_config["entries"]:
        dataset, metric, _response_metric, predictor, resp = _correlation_entry(entry)
        section = sections.setdefault(
            metric, {"title": metric_labels.get(metric, metric), "blocks": []})
        rows = []
        for m in models:
            key = (predictor, dataset, metric, resp, m["model"])
            row = tab.get(key)
            if row is None:
                raise KeyError(f"No influence-diagnostics row for {key}")
            p0 = _num(row, "p0")
            worst2_p, trio_p = _num(row, "worst_pair_p"), _num(row, "worst_trio_p")
            stay, n_pairs = _num(row, "n_pairs_stay_sig"), _num(row, "n_pairs")
            n_trios = _num(row, "n_trios")
            rows.append({"cells": [
                Cell(m["label"]),
                Cell(fmt_p_starred(p0), p0 < 0.05),
                Cell(_flip_cell(row, "loso_subject_flips", "n_subjects")),
                Cell(_flip_cell(row, "loso_structure_flips", "n_structures")),
                Cell(_flip_cell(row, "loco_flips", "n_combos")),
                Cell(f"{int(stay or 0)} / {int(n_pairs or 0)} significant "
                     f"(worst p = {_sig2(worst2_p)})"),
                # Leave-3 is run only for fits that are significant at raw x, so a
                # dash here means "not applicable", not "not robust".
                Cell("—" if trio_p is None
                     else f"worst p = {_sig2(trio_p)} over {int(n_trios or 0)} subsets"),
                Cell(f"slope {_sig2(_num(row, 'slope_excl_top2'))}, "
                     f"{fmt_p(_num(row, 'p_excl_top2'))}"),
            ]})
        section["blocks"].append({
            "label": (f"{partition_labels.get(dataset, dataset.upper())}, "
                      f"{predictor_labels.get(predictor, predictor)}"),
            "rows": rows,
        })
    return headers, list(sections.values())


def build_depth_profile(summary_config: dict, base_dir: Path):
    """S5: depth profile of OFF involvement: intercept-only LME group means with
    95% CIs (lmer + subject-cluster bootstrap) from
    _output_depth_profile/depth_profile_group_summary.csv."""
    import csv

    csv_path = base_dir / summary_config["summary_csv"]
    by_col = {}
    with open(csv_path) as f:
        for row in csv.DictReader(f):
            by_col[row["col"]] = row
    headers = ["Depth metric", "Group mean", "95% CI (LME)",
               "95% CI (subject bootstrap)", "n (subj-struct pairs)", "n subjects"]
    blocks = []
    for m in summary_config["metrics"]:
        r = by_col.get(m["col"])
        if r is None:
            continue

        def g(k):
            v = r.get(k, "")
            try:
                return float(v)
            except (TypeError, ValueError):
                return None
        cells = [
            Cell(_sig2(g("mean"))),
            Cell(_ci_str(g("lmer_lo"), g("lmer_hi"))),
            Cell(_ci_str(g("boot_lo"), g("boot_hi"))),
            Cell(str(r.get("n_combos", ""))),
            Cell(str(r.get("n_subjects", ""))),
        ]
        blocks.append({"label": m["label"], "rows": [{"cells": cells}]})
    return headers, [{"title": summary_config.get("section_title", ""), "blocks": blocks}]


def _read_csv_rows(path: Path):
    import csv

    with open(path) as f:
        return list(csv.DictReader(f))


def _fnum(row, key):
    try:
        return float(row[key])
    except (KeyError, ValueError, TypeError):
        return None


def build_per_event_meta(summary_config: dict, base_dir: Path):
    """S2a: per-event OFF-property vs co-occurring delta: pooled Spearman r
    (DerSimonian-Laird meta) + 95% CI, from the cnpix_local_sleep export CSV."""
    rows = _read_csv_rows(base_dir / summary_config["data_csv"])
    part_order = {"All OFFs": 0, "Medium + Large": 1, "Small OFFs": 2}
    cond_order = {"NREM": 0, "Wake": 1}
    headers = ["OFF property", "Detection", "State", "Pooled Spearman r [95% CI]",
               "p", "I² (%)", "k (subj-struct pairs)"]
    blocks: "OrderedDict[str, dict]" = OrderedDict()
    for r in rows:
        prop = r["property"]
        block = blocks.setdefault(prop, {"label": prop, "rows": []})
        p = _fnum(r, "p_value")
        pcell = fmt_p_starred(p) if p is not None else ""
        if r.get("sig_bonferroni", "").strip().lower() in ("false", "0", "no"):
            pcell += " (n.s. Bonf.)"
        i2 = _fnum(r, "i_squared")
        cells = [
            Cell(r["partition"]),
            Cell(r["condition"]),
            Cell(f'{_sig2(_fnum(r, "pooled_rho"))} {_ci_str(_fnum(r, "ci_lo"), _fnum(r, "ci_hi"))}',
                 p is not None and p < 0.05),
            Cell(pcell, p is not None and p < 0.05),
            Cell(f"{i2:.0f}" if i2 is not None else ""),
            Cell(str(r.get("k", ""))),
        ]
        sort = (part_order.get(r["partition"], 9), cond_order.get(r["condition"], 9))
        block["rows"].append({"cells": cells, "sort": sort})
    for b in blocks.values():
        b["rows"].sort(key=lambda x: x["sort"])
        for x in b["rows"]:
            del x["sort"]
    return headers, [{"title": summary_config.get("section_title", ""), "blocks": list(blocks.values())}]


def build_size_globality(summary_config: dict, base_dir: Path):
    """S3b (Fig. 4d-f): pooled Spearman of OFF size vs raw overlap degree and vs
    excess globality, from the cnpix_local_sleep export CSV."""
    rows = _read_csv_rows(base_dir / summary_config["data_csv"])
    prop_order = {"duration": 0, "span": 1, "area": 2}
    target_order = {"raw degree": 0, "excess": 1}
    prop_labels = {"duration": "OFF duration", "span": "OFF span", "area": "OFF area"}
    headers = ["OFF property", "Overlap target", "Pooled Spearman r [95% CI]", "p",
               "I² (%)", "k (subj-struct pairs)"]
    blocks: "OrderedDict[str, dict]" = OrderedDict()
    for r in sorted(rows, key=lambda x: (prop_order.get(x["property"], 9),
                                         target_order.get(x["target"], 9))):
        label = prop_labels.get(r["property"], r["property"])
        block = blocks.setdefault(label, {"label": label, "rows": []})
        p = _fnum(r, "p_value")
        i2 = _fnum(r, "i_squared")
        block["rows"].append({"cells": [
            Cell(r["target"]),
            Cell(f'{_sig2(_fnum(r, "pooled_rho"))} {_ci_str(_fnum(r, "ci_lo"), _fnum(r, "ci_hi"))}',
                 p is not None and p < 0.05),
            Cell(fmt_p_starred(p) if p is not None else "", p is not None and p < 0.05),
            Cell(f"{i2:.0f}" if i2 is not None else ""),
            Cell(str(r.get("k", ""))),
        ]})
    return headers, [{"title": summary_config.get("section_title", ""), "blocks": list(blocks.values())}]


def build_epoched_partial(summary_config: dict, base_dir: Path):
    """S2b: epoched partial: OFF area at each size level predicting epoch delta
    power; pooled standardized beta (marginal & partial) + 95% CI + I²."""
    rows = _read_csv_rows(base_dir / summary_config["data_csv"])
    part_order = {"Medium + Large": 0, "Small OFFs": 1}
    cond_order = {"NREM": 0, "Wake": 1}
    # semipartial is the reported partial quantity. The joint coefficient is not
    # tabulated: it is a rescaling of the same estimate testing the same null.
    model_order = {"marginal": 0, "semipartial": 1}
    model_labels = {"marginal": "marginal", "semipartial": "partial (semipartial)"}
    squareable = {"marginal", "semipartial"}  # models whose beta**2 is a variance share
    headers = ["Size level", "State", "Model", "Pooled std. β [95% CI]", "(pooled β)²",
               "p", "I² (%)", "k (subj-struct pairs)"]
    blocks: "OrderedDict[str, dict]" = OrderedDict()
    for r in rows:
        part = r["partition"]
        block = blocks.setdefault(part, {"label": part, "rows": []})
        p = _fnum(r, "p")
        i2 = _fnum(r, "i_squared")
        # Variance explained: the marginal beta squares to that level's R2, the
        # semipartial to its increment over the other level. Guarded by `squareable`
        # so a joint coefficient can never be squared into a bogus variance share.
        beta = _fnum(r, "pooled_std_beta")
        sq = f"{beta ** 2:.3f}" if beta is not None and r["model"] in squareable else ""
        cells = [
            Cell(r["condition"]),
            Cell(model_labels.get(r["model"], r["model"])),
            Cell(f'{_sig2(_fnum(r, "pooled_std_beta"))} {_ci_str(_fnum(r, "ci_lo"), _fnum(r, "ci_hi"))}',
                 p is not None and p < 0.05),
            Cell(sq),
            Cell(fmt_p_starred(p) if p is not None else "", p is not None and p < 0.05),
            Cell(f"{i2:.0f}" if i2 is not None else ""),
            Cell(str(r.get("k", ""))),
        ]
        sort = (cond_order.get(r["condition"], 9), model_order.get(r["model"], 9))
        block["rows"].append({"cells": cells, "sort": sort})
    for b in blocks.values():
        b["rows"].sort(key=lambda x: x["sort"])
        for x in b["rows"]:
            del x["sort"]
    order = sorted(blocks.values(), key=lambda b: part_order.get(b["label"], 9))
    return headers, [{"title": summary_config.get("section_title", ""), "blocks": order}]


def _read_varcomp_sds(leaf: Path):
    """Full-model random-effect SDs (sdcor) keyed by grouping factor."""
    out = {}
    vc = leaf / "variance_components.csv"
    if vc.exists():
        for row in _read_csv_rows(vc):
            if row.get("model") == "full":
                try:
                    out[row["grp"]] = float(row["sdcor"])
                except (ValueError, KeyError):
                    pass
    return out


def _n_obs(leaf: Path):
    ad = leaf / "adjusted_data.csv"
    if not ad.exists():
        return ""
    with open(ad) as f:
        return max(sum(1 for _ in f) - 1, 0)


def build_model_reference(summary_config: dict, base_dir: Path):
    """S0: per-model reference: response transform, inverse-variance weighting, n,
    and the fitted random-effect standard deviations, for every fit in the S1a
    homeostasis config (reused via ``from_config``)."""
    from_cfg = base_dir / summary_config["from_config"]
    with open(from_cfg) as f:
        src = yaml.safe_load(f)
    main_config = yaml.safe_load(open(base_dir / src["config"]))
    titles = load_rv_titles(main_config)
    partition_labels = {**DEFAULT_PARTITIONS, **(src.get("partition_labels") or {}),
                        "bandpower": "Cortical LFP"}
    set_labels = {"six": "6-cond", "nrem": "NREM", "wake": "Wake",
                  **(src.get("set_labels") or {})}
    prop_over = src.get("property_labels") or {}
    headers = ["Measure", "Detection", "Model set", "Transform", "Wt.", "n",
               "SD subj", "SD struct", "SD subj×struct", "SD resid"]

    out_sections = []
    for section in _section_configs(src):
        sec_results = base_dir / _resolve(section, src, "results_dir", "_output")
        blocks = []
        for rv, ds, cset, model in section["entries"]:
            leaf = sec_results / ds / cset / rv / model
            res = json.load(open(leaf / "results.json"))
            md = res.get("model_def", {})
            sds = _read_varcomp_sds(leaf)
            label = prop_over.get(rv) or titles["_generic"].get(rv, rv)
            cells = [
                Cell(partition_labels.get(ds, ds.upper())),
                Cell(set_labels.get(cset, cset)),
                Cell(res.get("transform", md.get("transform", "identity")) or "identity"),
                Cell("yes" if md.get("weighted") else "no"),
                Cell(str(_n_obs(leaf))),
                Cell(_sig2(sds.get("subject"))),
                Cell(_sig2(sds.get("structure"))),
                Cell(_sig2(sds.get("subject:structure"))),
                Cell(_sig2(sds.get("Residual"))),
            ]
            blocks.append({"label": label, "rows": [{"cells": cells}]})
        out_sections.append({"title": section.get("title", ""), "blocks": blocks})
    return headers, out_sections


CONDITION_SHORT = {
    "Early.REC.NREM.Match": "late baseline",
    "Early.REC.NREM": "early recovery",
    "Early.BSL.NREM": "early baseline",
    "Late.REC.NREM": "late recovery",
    "Early.NOD.Wake": "early SD",
    "Late.NOD.Wake": "late SD",
}


def prettify_contrast(s: str) -> str:
    """Replace condition tokens with short plain-language names (longest first)."""
    for k in sorted(CONDITION_SHORT, key=len, reverse=True):
        s = s.replace(k, CONDITION_SHORT[k])
    return s


def build_locality(summary_config: dict, base_dir: Path):
    """S3: OFF locality/globality, mirroring Fig 4. Contrasts are rows so that
    main-effect fits (overlap degree ~ condition [Fig 4b]; observed vs null
    [Fig 4c]) and interaction fits (overlapping - local size by condition, with the
    difference-in-differences [Fig 4g]) share one layout. Each ``fits`` entry:
    {results_dir, path, label, kind: main_effect|interaction, blocks: [...], units}.
    """
    headers = ["Analysis", "Units", "Contrast", "Omnibus effect",
               "Post-hoc (p, estimate [95% CI], d)"]
    out = []
    for section in summary_config["sections"]:
        blocks = []
        for fit in section["fits"]:
            results_dir = base_dir / _resolve(fit, summary_config, "results_dir", "_output_locality")
            res = json.load(open(results_dir / fit["path"] / "results.json"))
            units = fit.get("units", "")
            rows = []
            if fit.get("kind", "main_effect") == "interaction":
                it = res["interaction"]
                omnibus = fmt_main_effect(
                    {"pval": it["pval"], "significant": it.get("significant"),
                     "cohens_f2": it.get("cohens_f2")}
                )
                for blk in fit.get("blocks", ["simple", "did"]):
                    ph = it.get(blk)
                    if not isinstance(ph, dict) or not ph.get("contrasts"):
                        continue
                    for i, c in enumerate(ph["contrasts"]):
                        tag = "gap Δ: " if blk == "did" else ""
                        rows.append((tag + prettify_contrast(c), fmt_contrast(ph, i),
                                     ph["pvalues"][i] < 0.05))
            else:
                me = res["main_effect"]
                omnibus = fmt_main_effect(me)
                ph = me.get("posthoc") if me.get("significant") else None
                if ph and ph.get("contrasts"):
                    for i, c in enumerate(ph["contrasts"]):
                        rows.append((prettify_contrast(c), fmt_contrast(ph, i),
                                     ph["pvalues"][i] < 0.05))
            block_rows = []
            for j, (contrast, cell, sig) in enumerate(rows):
                block_rows.append({"cells": [
                    Cell(units if j == 0 else ""),
                    Cell(contrast),
                    Cell(omnibus if j == 0 else "", res.get("main_effect", {}).get("significant", False) if j == 0 else False),
                    Cell(cell, sig),
                ]})
            blocks.append({"label": fit["label"], "rows": block_rows or [{"cells": [Cell(units), Cell("—"), Cell(omnibus), Cell("—")]}]})
        out.append({"title": section.get("title", ""), "blocks": blocks})
    return headers, out


# Every builder now has the same (config, base_dir) signature.
TABLE_BUILDERS = {
    "model_reference": build_model_reference,
    "homeostasis": build_document,
    "locality": build_locality,
    "correlation_slope": build_correlation_slope,
    "correlation_robustness": build_correlation_robustness,
    "depth_profile": build_depth_profile,
    "per_event_meta": build_per_event_meta,
    "epoched_partial": build_epoched_partial,
    "size_globality": build_size_globality,
}


# -------------------- Emitters --------------------
def _iter_render_rows(sections):
    """Yield (kind, cells) render rows across all sections:
    ("section", [title]) | ("data", [property_or_blank, *cell_texts], sig_flags)."""
    for section in sections:
        if section.get("title"):
            yield ("section", section["title"], None)
        for block in section["blocks"]:
            for i, row in enumerate(block["rows"]):
                label = block["label"] if i == 0 else ""
                texts = [label] + [c.text for c in row["cells"]]
                sig = [False] + [c.significant for c in row["cells"]]
                yield ("data", texts, sig)


def _flat_rows(headers, sections):
    """Flatten to a plain list-of-lists (for CSV / console), section titles inline."""
    out = [list(headers)]
    for kind, cells, _sig in _iter_render_rows(sections):
        out.append(cells if kind == "data" else [cells])
    return out


def write_csv(headers, sections, path: Path):
    import csv

    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", newline="") as f:
        w = csv.writer(f)
        w.writerows(_flat_rows(headers, sections))


def _fill_xlsx_sheet(ws, headers, sections, title, caption):
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    thin = Side(style="thin", color="999999")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    section_fill = PatternFill("solid", fgColor="E8E8E8")
    full_headers = list(headers)
    ncol = len(full_headers)

    ws.cell(1, 1, title).font = Font(bold=True, size=12)
    for c, h in enumerate(full_headers, start=1):
        cell = ws.cell(2, c, h)
        cell.font = Font(bold=True)
        cell.alignment = center
        cell.border = Border(bottom=thin, top=thin)
    r = 3

    for section in sections:
        if section.get("title"):
            sc = ws.cell(r, 1, section["title"])
            sc.font = Font(bold=True, italic=True)
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=ncol)
            for c in range(1, ncol + 1):
                ws.cell(r, c).fill = section_fill
            r += 1
        for block in section["blocks"]:
            block_start = r
            for i, row in enumerate(block["rows"]):
                ws.cell(r, 1, block["label"] if i == 0 else None)
                for c, cellobj in enumerate(row["cells"], start=2):
                    x = ws.cell(r, c, cellobj.text)
                    x.alignment = left if c <= 3 else center
                    if cellobj.significant:
                        x.font = Font(bold=True)
                r += 1
            if r - 1 > block_start:
                ws.merge_cells(start_row=block_start, start_column=1, end_row=r - 1, end_column=1)
            ws.cell(block_start, 1).font = Font(bold=True)
            ws.cell(block_start, 1).alignment = Alignment(vertical="center", wrap_text=True)

    r += 1
    cap = ws.cell(r, 1, caption)
    cap.alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=ncol)

    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 8
    ws.column_dimensions["D"].width = 20
    for i in range(ncol - 4):
        ws.column_dimensions[chr(ord("E") + i)].width = 26
    ws.freeze_panes = ws.cell(3, 1)


def write_xlsx(headers, sections, path: Path, title: str, caption: str):
    from openpyxl import Workbook

    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    _fill_xlsx_sheet(wb.active, headers, sections, title, caption)
    wb.active.title = "table"
    wb.save(path)


def write_supplement_xlsx(tables, path: Path):
    """One workbook, one sheet per table. ``tables`` = [{sheet, headers, sections,
    title, caption}]."""
    from openpyxl import Workbook

    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)
    for t in tables:
        ws = wb.create_sheet(title=t["sheet"][:31])
        _fill_xlsx_sheet(ws, t["headers"], t["sections"], t["title"], t["caption"])
    wb.save(path)


def _import_docx():
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Pt

        return Document, WD_TABLE_ALIGNMENT, Pt
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(
            "python-docx is not installed. It is a declared dependency of "
            "cnpix-local-sleep; install that package and retry."
        ) from exc


def _append_docx_table(doc, headers, sections, title, caption):
    _, WD_TABLE_ALIGNMENT, Pt = _import_docx()
    doc.add_paragraph(title).runs[0].bold = True

    full_headers = list(headers)
    ncol = len(full_headers)
    table = doc.add_table(rows=1, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(full_headers):
        cell = table.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        lines = str(h).split("\n")
        for i, line in enumerate(lines):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(8)
            if i < len(lines) - 1:
                run.add_break()

    for kind, cells, sig in _iter_render_rows(sections):
        row = table.add_row()
        if kind == "section":
            merged = row.cells[0].merge(row.cells[ncol - 1])
            merged.text = cells if isinstance(cells, str) else cells[0]
            for p in merged.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.italic = True
                    run.font.size = Pt(8)
            continue
        for c, val in enumerate(cells):
            row.cells[c].text = val
            for p in row.cells[c].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    if c < len(sig) and sig[c]:
                        run.bold = True

    cap = doc.add_paragraph()
    cap.add_run(caption).font.size = Pt(8)


def write_docx(headers, sections, path: Path, title: str, caption: str):
    Document, _, _ = _import_docx()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _append_docx_table(doc, headers, sections, title, caption)
    doc.save(path)


def write_supplement_docx(tables, path: Path):
    """One document, all tables in order (page break between). ``tables`` =
    [{headers, sections, title, caption}]."""
    Document, _, _ = _import_docx()
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for i, t in enumerate(tables):
        if i > 0:
            doc.add_page_break()
        _append_docx_table(doc, t["headers"], t["sections"], t["title"], t["caption"])
    doc.save(path)


INHERITABLE_KEYS = ("entries", "partition_labels", "metric_labels", "predictor_labels")


def _inherit_entries(summary_config: dict, summary_yaml: Path) -> None:
    """Resolve ``entries_from: <sibling config>`` in place, so two tables that must
    cover exactly the same fits (S4a slopes / S4b robustness) declare that fit list
    once. Keys already present in this config win."""
    src = summary_config.get("entries_from")
    if not src:
        return
    with open((summary_yaml.resolve().parent / src)) as f:
        parent = yaml.safe_load(f)
    for key in INHERITABLE_KEYS:
        if key not in summary_config and key in parent:
            summary_config[key] = parent[key]


def render_config(summary_yaml: Path) -> dict:
    """Load one table config and build it. Returns a dict with headers, sections,
    title, caption, stem, sheet: the unit both single-table and supplement output
    consume."""
    with open(summary_yaml) as f:
        summary_config = yaml.safe_load(f)
    _inherit_entries(summary_config, summary_yaml)
    base_dir = resolve_base_dir(summary_yaml, summary_config)
    table_type = summary_config.get("table_type", "homeostasis")
    builder = TABLE_BUILDERS.get(table_type)
    if builder is None:
        raise SystemExit(f"Unknown table_type: {table_type!r} (have {list(TABLE_BUILDERS)})")
    headers, sections = builder(summary_config, base_dir)
    return {
        "headers": headers,
        "sections": sections,
        "title": summary_config.get("title", summary_yaml.stem),
        "caption": summary_config.get("caption", ""),
        "stem": summary_yaml.stem,
        "sheet": summary_config.get("sheet", summary_yaml.stem.replace("manuscript_", "")),
        "base_dir": base_dir,
        "output_dir": summary_config.get("output_dir", "_output_manuscript"),
    }


def main(argv=None):
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("summary_yaml", type=Path,
                        help="A table config, OR (with --supplement) a supplement config listing 'tables:'.")
    parser.add_argument("--supplement", action="store_true",
                        help="Treat summary_yaml as a supplement config (key 'tables:' = ordered list "
                             "of table-config paths) and write ONE combined .docx + .xlsx.")
    parser.add_argument("--csv", type=Path, default=None)
    parser.add_argument("--xlsx", type=Path, default=None)
    parser.add_argument("--docx", type=Path, default=None)
    args = parser.parse_args(argv)

    if args.supplement:
        with open(args.summary_yaml) as f:
            supp = yaml.safe_load(f)
        base = args.summary_yaml.resolve().parent
        tables = []
        for p in supp["tables"]:
            try:
                tables.append(render_config((base / p).resolve()))
            except (FileNotFoundError, KeyError) as e:
                print(f"  WARNING: skipping {p} (data not ready?): {e}", file=sys.stderr)
        out_dir = tables[0]["base_dir"] / supp.get("output_dir", "_output_manuscript")
        stem = args.summary_yaml.stem
        docx_path = args.docx or (out_dir / f"{stem}.docx")
        xlsx_path = args.xlsx or (out_dir / f"{stem}.xlsx")
        write_supplement_docx(tables, docx_path)
        write_supplement_xlsx(tables, xlsx_path)
        print(f"Supplement ({len(tables)} tables):\n  DOCX: {docx_path}\n  XLSX: {xlsx_path}",
              file=sys.stderr)
        return

    t = render_config(args.summary_yaml)
    for row in _flat_rows(t["headers"], t["sections"]):
        print(" | ".join(str(x) for x in row))
    out_dir = t["base_dir"] / t["output_dir"]
    csv_path = args.csv or (out_dir / f"{t['stem']}.csv")
    xlsx_path = args.xlsx or (out_dir / f"{t['stem']}.xlsx")
    write_csv(t["headers"], t["sections"], csv_path)
    write_xlsx(t["headers"], t["sections"], xlsx_path, t["title"], t["caption"])
    print(f"\nCSV : {csv_path}\nXLSX: {xlsx_path}", file=sys.stderr)
    docx_path = args.docx or (out_dir / f"{t['stem']}.docx")
    write_docx(t["headers"], t["sections"], docx_path, t["title"], t["caption"])
    print(f"DOCX: {docx_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
